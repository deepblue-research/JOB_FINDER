import os
import json
from docx import Document
from fastapi import APIRouter, HTTPException
from pydantic import BaseModel
import asyncpg

from services.gemini_service import generate_l3_round1, generate_l3_round2
from services.resume_formatter import format_resume

router = APIRouter()


class OptimizeRequest(BaseModel):
    session_id: str
    job_description: str


class OptimizeEnhanceRequest(BaseModel):
    session_id: str
    followup_answers: dict


def extract_resume_text(file_path: str) -> str:
    """Pull plain text out of a generated resume .docx for use as Gemini input."""
    doc = Document(file_path)
    lines = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    lines.append(cell.text)
    return "\n".join(lines)


@router.post("/")
async def optimize(request: OptimizeRequest):
    """Round 1: rewrite resume to match a job description and surface follow-up questions."""
    try:
        conn = await asyncpg.connect(os.getenv("ENRICHMENT_DATABASE_URL"))
        row = await conn.fetchrow(
            "SELECT resume_path, improved_resume_path FROM enrichment_sessions WHERE session_id = $1",
            request.session_id
        )

        if not row:
            await conn.close()
            raise HTTPException(status_code=404, detail="Session not found")

        # Use the most recent resume version (Level 2 if it exists, else Level 1)
        resume_path = row["improved_resume_path"] or row["resume_path"]

        if not resume_path or not os.path.exists(resume_path):
            await conn.close()
            raise HTTPException(status_code=400, detail="No resume found for this session. Complete Level 1 first.")

        resume_text = extract_resume_text(resume_path)

        # Call Gemini Round 1
        result = generate_l3_round1(resume_text, request.job_description)

        rewritten_resume = result["rewritten_resume"]
        followup_questions = result.get("followup_questions", [])
        skill_gaps = result.get("skill_gaps", [])
        jd_match_score = result.get("jd_match_score")

        # Save rewritten resume as .docx
        output_dir = os.path.join(os.path.dirname(__file__), "..", "outputs")
        output_path = os.path.join(output_dir, f"{request.session_id}_rewritten.docx")
        format_resume(rewritten_resume, output_path)

        # Save to database
        await conn.execute(
            """
            UPDATE enrichment_sessions
            SET jd_text = $1,
                l3_rewritten_resume = $2::jsonb,
                l3_followup_questions = $3::jsonb,
                skill_gaps = $4::jsonb,
                jd_match_score = $5,
                level_completed = 3,
                updated_at = NOW()
            WHERE session_id = $6
            """,
            request.job_description,
            json.dumps(rewritten_resume),
            json.dumps(followup_questions),
            json.dumps(skill_gaps),
            jd_match_score,
            request.session_id
        )
        await conn.close()

        return {
            "message": "Resume rewritten for this job. Answer the follow-up questions to finalize.",
            "rewritten_download_url": f"/optimize/download/{request.session_id}",
            "followup_questions": followup_questions,
            "skill_gaps": skill_gaps,
            "jd_match_score": jd_match_score
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/enhance")
async def optimize_enhance(request: OptimizeEnhanceRequest):
    """Round 2: incorporate follow-up answers and produce the final tailored resume."""
    try:
        conn = await asyncpg.connect(os.getenv("ENRICHMENT_DATABASE_URL"))
        row = await conn.fetchrow(
            "SELECT l3_rewritten_resume, jd_text, skill_gaps FROM enrichment_sessions WHERE session_id = $1",
            request.session_id
        )

        if not row or not row["l3_rewritten_resume"]:
            await conn.close()
            raise HTTPException(status_code=404, detail="No Round 1 result found. Call /optimize first.")

        rewritten_resume = json.loads(row["l3_rewritten_resume"])
        jd_text = row["jd_text"]
        skill_gaps = json.loads(row["skill_gaps"]) if row["skill_gaps"] else []

        # Call Gemini Round 2
        result = generate_l3_round2(rewritten_resume, request.followup_answers, jd_text)

        final_resume = result["final_resume"]
        changes_summary = result.get("changes_summary", [])

        # Save final resume as .docx
        output_dir = os.path.join(os.path.dirname(__file__), "..", "outputs")
        output_path = os.path.join(output_dir, f"{request.session_id}_final.docx")
        format_resume(final_resume, output_path)

        # Save to database
        await conn.execute(
            """
            UPDATE enrichment_sessions
            SET l3_followup_answers = $1::jsonb,
                final_resume_path = $2,
                updated_at = NOW()
            WHERE session_id = $3
            """,
            json.dumps(request.followup_answers),
            output_path,
            request.session_id
        )
        await conn.close()

        return {
            "message": "Final tailored resume ready.",
            "download_url": f"/optimize/download-final/{request.session_id}",
            "changes_summary": changes_summary,
            "skill_gaps": skill_gaps
        }

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/download/{session_id}")
async def download_rewritten(session_id: str):
    """Download the Round 1 rewritten resume."""
    output_dir = os.path.join(os.path.dirname(__file__), "..", "outputs")
    file_path = os.path.join(output_dir, f"{session_id}_rewritten.docx")

    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Rewritten resume not found")

    from fastapi.responses import FileResponse
    return FileResponse(
        path=file_path,
        filename="resume_tailored_draft.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@router.get("/download-final/{session_id}")
async def download_final(session_id: str):
    """Download the final tailored resume after Round 2."""
    output_dir = os.path.join(os.path.dirname(__file__), "..", "outputs")
    file_path = os.path.join(output_dir, f"{session_id}_final.docx")

    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Final resume not found")

    from fastapi.responses import FileResponse
    return FileResponse(
        path=file_path,
        filename="resume_tailored_final.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@router.get("/skill-gaps/{session_id}")
async def get_skill_gaps(session_id: str):
    """Return skill gaps and course links from the last optimization."""
    conn = await asyncpg.connect(os.getenv("ENRICHMENT_DATABASE_URL"))
    row = await conn.fetchrow(
        "SELECT skill_gaps, jd_match_score FROM enrichment_sessions WHERE session_id = $1",
        session_id
    )
    await conn.close()

    if not row or not row["skill_gaps"]:
        raise HTTPException(status_code=404, detail="No skill gap data found for this session")

    return {
        "skill_gaps": json.loads(row["skill_gaps"]),
        "jd_match_score": row["jd_match_score"]
    }
