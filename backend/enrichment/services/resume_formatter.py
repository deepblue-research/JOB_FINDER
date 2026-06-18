from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

INDIGO = RGBColor(0x3D, 0x52, 0xA0)

def add_section_heading(doc, title):
    para = doc.add_paragraph()
    run = para.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = INDIGO
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(2)
    # Add a horizontal line under the heading
    border = para._p.get_or_add_pPr()
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '3D52A0')
    pBdr.append(bottom)
    border.append(pBdr)

def add_bullet(doc, text):
    para = doc.add_paragraph(style='List Bullet')
    run = para.add_run(text)
    run.font.size = Pt(10)
    para.paragraph_format.space_after = Pt(1)

def add_normal(doc, text, bold=False, size=10):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    para.paragraph_format.space_after = Pt(1)
    return para

def format_resume(resume_json: dict, output_path: str):
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # --- CONTACT INFO ---
    contact = resume_json.get("contact_info", {})
    if contact:
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(contact.get("name", ""))
        name_run.bold = True
        name_run.font.size = Pt(16)
        name_run.font.color.rgb = INDIGO
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        details = []
        if contact.get("phone"): details.append(contact["phone"])
        if contact.get("email"): details.append(contact["email"])
        if contact.get("location"): details.append(contact["location"])
        if contact.get("linkedin"): details.append(contact["linkedin"])
        if contact.get("github"): details.append(contact["github"])

        if details:
            detail_para = doc.add_paragraph(" | ".join(details))
            detail_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in detail_para.runs:
                run.font.size = Pt(9)

    # --- EDUCATION ---
    education = resume_json.get("education", {})
    if isinstance(education, list):
        education = education[0] if education else {}
    if education:
        add_section_heading(doc, "Education")
        degree_line = f"{education.get('degree', '')} in {education.get('branch', '')} — {education.get('college', '')}"
        p = add_normal(doc, degree_line, bold=True)
        if education.get("graduation_year"):
            add_normal(doc, f"Expected Graduation: {education['graduation_year']}")
        if education.get("cgpa"):
            add_normal(doc, f"CGPA: {education['cgpa']}")
        if education.get("class_12"):
            add_normal(doc, f"Class 12: {education['class_12']}")

    # --- SKILLS ---
    skills = resume_json.get("skills", {})
    if skills:
        add_section_heading(doc, "Skills")
        if isinstance(skills, list):
            # Gemini sometimes returns a flat list instead of a dict — handle both
            add_normal(doc, ", ".join(str(s) for s in skills))
        elif isinstance(skills, dict):
            if skills.get("programming_languages"):
                add_normal(doc, f"Languages: {', '.join(skills['programming_languages'])}")
            if skills.get("web_frameworks"):
                add_normal(doc, f"Frameworks: {', '.join(skills['web_frameworks'])}")
            if skills.get("frameworks"):
                add_normal(doc, f"Frameworks: {', '.join(skills['frameworks'])}")
            if skills.get("tools_and_technologies"):
                add_normal(doc, f"Tools: {', '.join(skills['tools_and_technologies'])}")
            if skills.get("tools"):
                add_normal(doc, f"Tools: {', '.join(skills['tools'])}")
            if skills.get("apis"):
                add_normal(doc, f"APIs: {', '.join(skills['apis'])}")
            if skills.get("soft_skills"):
                add_normal(doc, f"Soft Skills: {skills['soft_skills']}")
        else:
            add_normal(doc, str(skills))

    # --- PROJECTS ---
    projects = resume_json.get("projects", [])
    if projects:
        add_section_heading(doc, "Projects")
        for project in projects:
            add_normal(doc, project.get("name", ""), bold=True)
            if project.get("technologies"):
                add_normal(doc, f"Technologies: {project['technologies']}")
            # Gemini sometimes uses "bullets" instead of "description" — check both
            description = project.get("description") or project.get("bullets")
            if description:
                if isinstance(description, list):
                    for bullet in description:
                        add_bullet(doc, bullet)
                else:
                    add_bullet(doc, description)
            if project.get("github_url"):
                add_normal(doc, f"Link: {project['github_url']}")
            if project.get("award"):
                add_normal(doc, f"Recognition: {project['award']}")

    # --- RESEARCH ---
    research = resume_json.get("research", {})
    if isinstance(research, list):
        # Gemini sometimes returns a list of research entries — handle that too
        research = research[0] if research else {}
    if isinstance(research, dict) and research.get("topic"):
        add_section_heading(doc, "Research")
        if research.get("professor") and research.get("institution"):
            add_normal(doc, f"Under {research['professor']}, {research['institution']}", bold=True)
        if research.get("topic"):
            add_normal(doc, f"Topic: {research['topic']}")
        if research.get("duration"):
            add_normal(doc, f"Duration: {research['duration']}")
        if research.get("output"):
            add_normal(doc, f"Output: {research['output']}")
        if research.get("tools"):
            add_normal(doc, f"Tools: {research['tools']}")

    # --- INTERNSHIPS ---
    internships = resume_json.get("internships", [])
    if internships:
        add_section_heading(doc, "Internships")
        if isinstance(internships, list):
            for internship in internships:
                add_normal(doc, f"{internship.get('role', '')} — {internship.get('company', '')}", bold=True)
                if internship.get("duration"):
                    add_normal(doc, internship["duration"])
                if internship.get("tasks"):
                    if isinstance(internship["tasks"], list):
                        for task in internship["tasks"]:
                            add_bullet(doc, task)
                    else:
                        add_bullet(doc, internship["tasks"])
        else:
            add_normal(doc, f"{internships.get('role', '')} — {internships.get('company', '')}", bold=True)
            if internships.get("duration"):
                add_normal(doc, internships["duration"])

    # --- WORK EXPERIENCE ---
    work = resume_json.get("work_experience", [])
    if work:
        add_section_heading(doc, "Work Experience")
        if isinstance(work, list):
            for job in work:
                add_normal(doc, f"{job.get('role', '')} — {job.get('company', '')}", bold=True)
                if job.get("duration"):
                    add_normal(doc, job["duration"])
                if job.get("description"):
                    description = job["description"]
                    if isinstance(description, list):
                        for bullet in description:
                            add_bullet(doc, bullet)
                    else:
                        add_bullet(doc, description)
        else:
            add_normal(doc, f"{work.get('role', '')} — {work.get('company', '')}", bold=True)
            
    # --- CERTIFICATIONS ---
    certifications = resume_json.get("certifications", [])
    if certifications:
        add_section_heading(doc, "Certifications")
        if isinstance(certifications, list):
            for cert in certifications:
                add_bullet(doc, cert)
        else:
            add_bullet(doc, str(certifications))

    # --- ACHIEVEMENTS ---
    achievements = resume_json.get("achievements", {})
    if achievements:
        add_section_heading(doc, "Achievements")
        if isinstance(achievements, list):
            # Gemini sometimes returns a flat list instead of a dict — handle both
            for item in achievements:
                add_bullet(doc, str(item))
        elif isinstance(achievements, dict):
            if achievements.get("competitions"):
                add_bullet(doc, achievements["competitions"])
            if achievements.get("scholarships"):
                add_bullet(doc, achievements["scholarships"])
            if achievements.get("rankings"):
                add_bullet(doc, achievements["rankings"])
            if achievements.get("publications"):
                add_bullet(doc, achievements["publications"])
        else:
            add_normal(doc, str(achievements))

    # --- EXTRACURRICULARS ---
    extra = resume_json.get("extracurriculars", {})
    if extra:
        add_section_heading(doc, "Extracurriculars")
        if isinstance(extra, list):
            # Gemini sometimes returns a flat list instead of a dict — handle both
            for item in extra:
                add_bullet(doc, str(item))
        elif isinstance(extra, dict):
            if extra.get("clubs"):
                add_normal(doc, extra["clubs"])
            if extra.get("events"):
                add_bullet(doc, extra["events"])
            if extra.get("creative_output"):
                add_bullet(doc, extra["creative_output"])
        else:
            add_normal(doc, str(extra))

    # --- LANGUAGES ---
    languages = resume_json.get("languages", [])
    if languages:
        add_section_heading(doc, "Languages")
        if isinstance(languages, list):
            add_normal(doc, ", ".join(languages))
        else:
            add_normal(doc, str(languages))

    # --- INTERESTS ---
    interests = resume_json.get("interests", {})
    if interests:
        add_section_heading(doc, "Interests")
        if isinstance(interests, list):
            # Gemini sometimes returns a flat list instead of a dict — handle both
            add_normal(doc, ", ".join(str(i) for i in interests))
        elif isinstance(interests, dict):
            if interests.get("tech_interests"):
                tech = interests["tech_interests"]
                if isinstance(tech, list):
                    tech = ", ".join(tech)
                add_normal(doc, f"Tech: {tech}")
            if interests.get("hobbies"):
                hobbies = interests["hobbies"]
                if isinstance(hobbies, list):
                    hobbies = ", ".join(hobbies)
                add_normal(doc, f"Hobbies: {hobbies}")
        else:
            add_normal(doc, str(interests))

    doc.save(output_path)